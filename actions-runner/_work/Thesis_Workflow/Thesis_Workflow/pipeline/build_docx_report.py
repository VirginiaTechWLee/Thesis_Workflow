"""
Build a DOCX report from pipeline markdown reports and Femap PNG images.

Reads all numbered markdown reports from REPORTS_DIR, converts them to
formatted DOCX paragraphs, and embeds any PNG images from an images
subdirectory. Produces a single professional DOCX with cover page and TOC.

Usage:
    python pipeline/build_docx_report.py --reports-dir <path> [--images-dir <path>] [--output <path>]

If --images-dir is omitted, looks for <reports-dir>/images/
If --output is omitted, writes to <reports-dir>/pipeline_report.docx

Requires:
    pip install python-docx

If python-docx is not installed, exits with code 0 and a warning — never blocks pipeline.
"""

import os
import re
import sys
import argparse
from datetime import datetime


def check_docx():
    """Check if python-docx is available."""
    try:
        import docx
        return True
    except ImportError:
        print("WARNING: python-docx not installed — skipping DOCX generation")
        print("  Install with: pip install python-docx")
        return False


def discover_reports(reports_dir):
    """Find all numbered report files (01_ through 99_) in order."""
    reports = []
    for fname in sorted(os.listdir(reports_dir)):
        if re.match(r"^\d{2}_.*\.md$", fname) and not fname.startswith("00_"):
            fpath = os.path.join(reports_dir, fname)
            reports.append((fname, fpath))
    return reports


def discover_images(images_dir):
    """Find all PNG images in order."""
    if not images_dir or not os.path.isdir(images_dir):
        return []
    images = []
    for fname in sorted(os.listdir(images_dir)):
        if fname.lower().endswith(".png"):
            images.append((fname, os.path.join(images_dir, fname)))
    return images


def md_to_docx_paragraphs(doc, content):
    """Convert markdown content to DOCX paragraphs (simplified parser)."""
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for line in content.splitlines():
        stripped = line.strip()

        # Skip horizontal rules
        if stripped in ("---", "***", "___"):
            doc.add_paragraph("_" * 60)
            continue

        # Headers
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:].strip(), level=1)
            continue
        if stripped.startswith("## "):
            p = doc.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:].strip(), level=3)
            continue

        # Bullet points
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, text)
            continue

        # Numbered lists
        m = re.match(r"^\d+\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_formatted_text(p, m.group(1))
            continue

        # Empty lines
        if not stripped:
            doc.add_paragraph("")
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, stripped)


def add_formatted_text(paragraph, text):
    """Add text with basic bold/italic formatting."""
    # Split on bold (**text**) and italic (*text*) markers
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def build_docx(reports_dir, images_dir, output_path):
    """Build the DOCX report."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # --- Cover page ---
    study_name = os.path.basename(reports_dir)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n")
    p = doc.add_heading("Pipeline Analysis Report", level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Study: {study_name}").bold = True
    p.add_run(f"\nGenerated: {timestamp}")
    p.add_run("\nSpacecraft Bolt Looseness Detection Pipeline")
    p.add_run("\nVirginia Tech — Wayne Lee")

    doc.add_page_break()

    # --- Table of Contents placeholder ---
    doc.add_heading("Table of Contents", level=1)
    reports = discover_reports(reports_dir)
    for fname, fpath in reports:
        with open(fpath, "r", encoding="utf-8") as f:
            content = f.read()
        # Extract title
        title = fname
        for line in content.splitlines():
            if line.strip().startswith("# "):
                title = line.strip()[2:].strip()
                break
        p = doc.add_paragraph(f"  {title}", style="List Bullet")

    images = discover_images(images_dir)
    if images:
        p = doc.add_paragraph("  Appendix: Femap Model Images", style="List Bullet")

    doc.add_page_break()

    # --- Report sections ---
    for fname, fpath in reports:
        with open(fpath, "r", encoding="utf-8") as f:
            content = f.read()
        md_to_docx_paragraphs(doc, content)
        doc.add_page_break()

    # --- Image appendix ---
    if images:
        doc.add_heading("Appendix: Femap Model Images", level=1)
        doc.add_paragraph(f"{len(images)} images captured from Femap COM automation.")
        doc.add_paragraph("")

        for fname, fpath in images:
            # Image title from filename
            title = fname.replace(".png", "").replace("_", " ").title()
            doc.add_heading(title, level=2)
            try:
                doc.add_picture(fpath, width=Inches(6.0))
            except Exception as e:
                doc.add_paragraph(f"[Image failed to embed: {fname} — {e}]")
            doc.add_paragraph("")

    # --- Footer ---
    p = doc.add_paragraph()
    p.add_run(f"\nEnd of report — {len(reports)} sections, {len(images)} images")
    p.add_run(f"\nGenerated {timestamp}")

    doc.save(output_path)
    print(f"DOCX report written: {output_path}")
    print(f"  Sections: {len(reports)}")
    print(f"  Images: {len(images)}")


def main():
    parser = argparse.ArgumentParser(description="Build DOCX report from pipeline reports and images")
    parser.add_argument("--reports-dir", required=True, help="Directory containing numbered markdown reports")
    parser.add_argument("--images-dir", help="Directory containing PNG images (default: <reports-dir>/images/)")
    parser.add_argument("--output", help="Output DOCX path (default: <reports-dir>/pipeline_report.docx)")
    args = parser.parse_args()

    if not check_docx():
        sys.exit(0)

    if not os.path.isdir(args.reports_dir):
        print(f"ERROR: Reports directory not found: {args.reports_dir}")
        sys.exit(1)

    images_dir = args.images_dir or os.path.join(args.reports_dir, "images")
    output_path = args.output or os.path.join(args.reports_dir, "pipeline_report.docx")

    build_docx(args.reports_dir, images_dir, output_path)


if __name__ == "__main__":
    main()
